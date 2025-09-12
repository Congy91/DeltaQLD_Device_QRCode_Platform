# app.py - QR Device Manager with Sites, LAN QR links, and login-on-scan flow
import os, io, socket, argparse
from datetime import datetime
from urllib.parse import urlparse
from flask import Flask, render_template, request, redirect, url_for, session, send_file, flash
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
import qrcode, pandas as pd
from docx import Document
from docx.shared import Inches
from PIL import Image

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-key-change-me')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///devices.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# cookie settings appropriate for local LAN (HTTP)
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = False

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# ---------------- Models ----------------
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    role = db.Column(db.String(32), default='user')
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Site(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), unique=True, nullable=False)
    url = db.Column(db.String(500), nullable=True)  # new custom URL field
    devices = db.relationship('Device', backref='site', lazy=True)


class Device(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    device_name = db.Column(db.String(200))
    device_number = db.Column(db.String(200))
    device_type = db.Column(db.String(200))
    device_serial = db.Column(db.String(200))
    device_ip = db.Column(db.String(200))
    device_mac = db.Column(db.String(50))       
    device_firmware = db.Column(db.String(50))   
    location = db.Column(db.String(200))         
    site_id = db.Column(db.Integer, db.ForeignKey('site.id'), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

# ---------------- Auth ----------------
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# ---------------- Helpers ----------------
def get_local_ip():
    """Return the machine's LAN IP address (works even with multiple adapters)."""
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        # doesn't actually connect to the internet, just picks the right NIC
        s.connect(('8.8.8.8', 80))
        ip = s.getsockname()[0]
    except Exception:
        ip = '127.0.0.1'
    finally:
        s.close()
    return ip

def safe_redirect_path(next_url):
    """Return a safe internal path extracted from next_url (avoid open redirects)."""
    if not next_url:
        return url_for('index')
    parsed = urlparse(next_url)
    return parsed.path or url_for('index')

@app.route('/edit_site/<int:site_id>', methods=['POST'])
@login_required
def edit_site(site_id):
    site = Site.query.get_or_404(site_id)
    name = request.form.get('site_name', '').strip()
    url = request.form.get('site_url', '').strip()
    if name:
        site.name = name
    if url:
        site.url = url
    db.session.commit()
    flash('Site details updated', 'success')
    return redirect(url_for('site_view', site_id=site.id))


# ---------------- Public pages (site + QR creation do not require login) ----------------
@app.route('/')
def index():
    sites = Site.query.order_by(Site.name).all()
    return render_template('index.html', sites=sites)

@app.route('/add_site', methods=['POST'])
@login_required
def add_site():
    site_name = request.form.get('site_name', '').strip()
    site_url = request.form.get('site_url', '').strip() or None

    if site_name:
        # Create the site
        s = Site(name=site_name, url=site_url)
        db.session.add(s)
        db.session.commit()

        # Create a folder for this site documents
        site_folder = os.path.join('static', 'sites', str(s.id))
        os.makedirs(site_folder, exist_ok=True)

        # Create 3 default documents
        default_docs = ["Functional Description", "Points List", "Network Architecture"]
        for doc_name in default_docs:
            doc_path = os.path.join(site_folder, f"{doc_name.replace(' ', '_')}.docx")
            # Create empty Word doc
            from docx import Document
            doc = Document()
            doc.add_heading(doc_name, level=1)
            doc.save(doc_path)

        flash(f'Site "{site_name}" created with default documents.', 'success')

    return redirect(url_for('index'))



@app.route('/add_device', methods=['POST'])
def add_device():
    # public endpoint to create a device record (useful before sticking QR)
    device_number = request.form.get('device_number','').strip()
    site_id = request.form.get('site_id') or None
    if not device_number:
        flash('Device number required','danger'); return redirect(url_for('index'))
    d = Device(device_number=device_number)
    if site_id:
        try: d.site_id = int(site_id)
        except: d.site_id = None
    db.session.add(d); db.session.commit()
    flash(f'Device created (ID {d.id})', 'success')
    return redirect(url_for('index'))

from docx import Document
from docx.shared import Inches

@app.route('/export_site_qrs/<int:site_id>')
@login_required
def export_site_qrs(site_id):
    site = Site.query.get_or_404(site_id)
    devices = Device.query.filter_by(site_id=site_id).all()

    doc = Document()
    doc.add_heading(f'Site: {site.name} - Device QR Codes', level=1)

    for d in devices:
        doc.add_paragraph(f'Device: {d.device_name or "(no name)"}')
        doc.add_paragraph(f'Number: {d.device_number or ""}')
        
        # generate QR image for this device
        import qrcode
        import io
        img = generate_qr_with_logo(f"http://{get_local_ip()}:5000/device/{d.id}?scan=1")
        bio = io.BytesIO()
        img.save(bio, 'PNG')
        bio.seek(0)
        
        doc.add_picture(bio, width=Inches(1.5))  # adjust size as needed
        doc.add_paragraph('')  # space after each device

    # Save to BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(output,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     download_name=f'{site.name}_QR_codes.docx',
                     as_attachment=True)
    
def generate_qr_with_logo(data, qr_size=300, logo_ratio=0.2):
    """
    Generates a QR code PIL image with a centered logo.
    Uses QRLogo.png from the static folder.
    """
    logo_path = os.path.join('static', 'QRLogo.png')
    
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,  # high error correction for logo
        box_size=10,
        border=4,
    )
    qr.add_data(data)
    qr.make(fit=True)

    img_qr = qr.make_image(fill_color="black", back_color="white").convert('RGB')
    img_qr = img_qr.resize((qr_size, qr_size))

    # Add logo
    logo = Image.open('static/QRLogo.png')
    logo_size = int(qr_size * logo_ratio)
    logo = logo.resize((logo_size, logo_size))
    pos = ((qr_size - logo_size) // 2, (qr_size - logo_size) // 2)
    img_qr.paste(logo, pos, mask=logo if logo.mode=='RGBA' else None)

    return img_qr


@app.route('/generate_qr/device/<int:device_id>')
def generate_device_qr(device_id):
    ip = get_local_ip()
    port = request.environ.get('SERVER_PORT', '5000')
    device_url = f'http://{ip}:{port}{url_for("device_view", device_id=device_id)}?scan=1'
    img = generate_qr_with_logo(device_url)
    bio = io.BytesIO(); img.save(bio, 'PNG'); bio.seek(0)
    return send_file(bio, mimetype='image/png', download_name=f'qr_device_{device_id}.png', as_attachment=True)



@app.route('/generate_qr/site/<int:site_id>')
def generate_site_qr(site_id):
    ip = get_local_ip()
    port = request.environ.get('SERVER_PORT', '5000')
    site_url = f'http://{ip}:{port}{url_for("site_view", site_id=site_id)}?scan=1'
    img = generate_qr_with_logo(site_url)
    bio = io.BytesIO(); img.save(bio, 'PNG'); bio.seek(0)
    return send_file(bio, mimetype='image/png', download_name=f'qr_site_{site_id}.png', as_attachment=True)



# ---------------- Protected pages (viewing/editing device or site requires login) ----------------
@app.route('/site/<int:site_id>', methods=['GET', 'POST'])
def site_view(site_id):
    # Force re-login if QR scan
    if 'scan' in request.args:
        session.pop('user_id', None)
        return redirect(url_for('login', next=url_for('site_view', site_id=site_id)))

    if not current_user.is_authenticated:
        return redirect(url_for('login', next=url_for('site_view', site_id=site_id)))

    site = Site.query.get_or_404(site_id)
    devices = Device.query.filter_by(site_id=site_id).all()

    if request.method == 'POST':
        # Update site name and URL
        new_name = request.form.get('site_name', '').strip()
        new_url = request.form.get('site_url', '').strip()
        if new_name:
            site.name = new_name
        site.url = new_url or None  # store None if empty
        db.session.commit()
        flash('Site updated successfully', 'success')
        return redirect(url_for('site_view', site_id=site.id))

    return render_template('site.html', site=site, devices=devices)


@app.route('/device/<int:device_id>', methods=['GET','POST'])
def device_view(device_id):
    device = Device.query.get_or_404(device_id)
    sites = Site.query.order_by(Site.name).all()
    logged_in = current_user.is_authenticated

    # Handle edits only if logged in
    if logged_in and request.method == 'POST':
        device.device_name = request.form.get('device_name') or device.device_name
        device.device_number = request.form.get('device_number') or device.device_number
        device.device_type = request.form.get('device_type') or device.device_type
        device.device_serial = request.form.get('device_serial') or device.device_serial
        device.device_ip = request.form.get('device_ip') or device.device_ip

        site_choice = request.form.get('site_choice')
        new_site = request.form.get('new_site','').strip()
        if new_site:
            s = Site.query.filter_by(name=new_site).first()
            if not s:
                s = Site(name=new_site)
                db.session.add(s)
                db.session.commit()
            device.site_id = s.id
        elif site_choice:
            try: device.site_id = int(site_choice)
            except: device.site_id = None

        db.session.add(device)
        db.session.commit()
        flash('Device saved','success')
        return redirect(url_for('device_view', device_id=device.id))

    return render_template('device.html', device=device, sites=sites, logged_in=logged_in)



@app.route('/export_site/<int:site_id>')
@login_required
def export_site(site_id):
    import openpyxl
    from openpyxl.styles import Font
    from openpyxl.utils import get_column_letter

    site = Site.query.get_or_404(site_id)
    devices = Device.query.filter_by(site_id=site_id).all()
    rows = []
    for d in devices:
        rows.append({
            'ID': d.id,
            'Device Name': d.device_name,
            'Device Number': d.device_number,
            'Device Type': d.device_type,
            'Serial Number': d.device_serial,
            'IP Address': d.device_ip,
            'MAC Address': d.device_mac,
            'Firmware': d.device_firmware,
            'Location': d.location
        })

    df = pd.DataFrame(rows)

    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Devices')
        ws = writer.sheets['Devices']

        # Bold headers
        for cell in ws[1]:
            cell.font = Font(bold=True)

        # Auto-width columns
        for col in ws.columns:
            max_length = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = max_length + 2

        # Add filter to header row
        ws.auto_filter.ref = ws.dimensions

        # Freeze top row
        ws.freeze_panes = ws['A2']

    bio.seek(0)
    return send_file(
        bio,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        download_name=f'{site.name}_devices.xlsx',
        as_attachment=True
    )

    
@app.route('/delete_site/<int:site_id>', methods=['POST'])
@login_required
def delete_site(site_id):
    site = Site.query.get_or_404(site_id)

    # Delete all devices for this site first
    Device.query.filter_by(site_id=site.id).delete()

    # Then delete the site itself
    db.session.delete(site)
    db.session.commit()
    flash(f'Site "{site.name}" and all associated devices have been deleted.', 'success')
    return redirect(url_for('index'))
    
@app.route('/delete_device/<int:device_id>', methods=['POST'])
@login_required
def delete_device(device_id):
    device = Device.query.get_or_404(device_id)
    site_id = device.site_id
    db.session.delete(device)
    db.session.commit()
    flash(f'Device {device.device_name or device.device_number} deleted', 'success')
    return redirect(url_for('site_view', site_id=site_id))

@app.route('/qr_image/device/<int:device_id>')
def qr_image_device(device_id):
    ip = get_local_ip()
    port = request.environ.get('SERVER_PORT', '5000')
    device_url = f'http://{ip}:{port}{url_for("device_view", device_id=device_id)}?scan=1'
    img = qrcode.make(device_url)
    bio = io.BytesIO()
    img.save(bio, 'PNG')
    bio.seek(0)
    return send_file(bio, mimetype='image/png')

@app.route('/qr_image/site/<int:site_id>')
def qr_image_site(site_id):
    ip = get_local_ip()
    port = request.environ.get('SERVER_PORT', '5000')
    site_url = f'http://{ip}:{port}{url_for("site_view", site_id=site_id)}?scan=1'
    img = qrcode.make(site_url)
    bio = io.BytesIO()
    img.save(bio, 'PNG')
    bio.seek(0)
    return send_file(bio, mimetype='image/png')


@app.route('/preview_qr/device/<int:device_id>')
def preview_qr_device(device_id):
    device_url = f'http://{get_local_ip()}:5000/device/{device_id}?scan=1'
    img = generate_qr_with_logo(device_url)
    bio = io.BytesIO()
    img.save(bio, 'PNG')
    bio.seek(0)
    return send_file(bio, mimetype='image/png')

@app.route('/preview_qr/site/<int:site_id>')
def preview_qr_site(site_id):
    site_url = f'http://{get_local_ip()}:5000/site/{site_id}?scan=1'
    img = generate_qr_with_logo(site_url)
    bio = io.BytesIO()
    img.save(bio, 'PNG')
    bio.seek(0)
    return send_file(bio, mimetype='image/png')


@app.route('/import_site', methods=['GET','POST'])
@login_required
def import_site():
    sites = Site.query.order_by(Site.name).all()
    if request.method == 'POST':
        f = request.files.get('excel_file')
        site_id = request.form.get('site_id') or None
        if not f:
            flash('No file uploaded','danger'); return redirect(url_for('import_site'))
        df = pd.read_excel(f)
        for idx, row in df.iterrows():
            d = None
            if 'id' in row and not pd.isna(row['id']): d = Device.query.get(int(row['id']))
            if d is None and 'device_number' in row and not pd.isna(row['device_number']):
                d = Device.query.filter_by(device_number=str(row['device_number'])).first()
            if d is None: d = Device()
            for col in ['device_name','device_number','device_type','device_serial','device_ip']:
                if col in row and not pd.isna(row[col]):
                    setattr(d, col, str(row[col]))
            # Assign site from dropdown
            if site_id:
                d.site_id = int(site_id)
            db.session.add(d)
        db.session.commit()
        flash('Import complete','success')
        return redirect(url_for('index'))
    return render_template('import.html', sites=sites)


# ---------------- Authentication ----------------
@app.route('/login', methods=['GET','POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username'); password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            login_user(user)
            next_url = request.args.get('next') or request.form.get('next')
            return redirect(safe_redirect_path(next_url))
        flash('Invalid credentials','danger')
    next_url = request.args.get('next') or request.form.get('next') or ''
    return render_template('login.html', next=next_url)

@app.route('/logout')
@login_required
def logout():
    logout_user(); flash('Logged out','info'); return redirect(url_for('index'))

# ---------------- CLI helpers ----------------
def init_db():
    with app.app_context():
        db.create_all()
        print('DB initialized.')

def create_admin(username, password):
    with app.app_context():
        if User.query.filter_by(username=username).first():
            print('User exists'); return
        u = User(username=username, role='admin'); u.set_password(password)
        db.session.add(u); db.session.commit()
        print('Admin created.')

# ---------------- Runner ----------------
if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('cmd', nargs='?', help='initdb | create-admin | run')
    parser.add_argument('--username', help='admin username')
    parser.add_argument('--password', help='admin password')
    args = parser.parse_args()
    if args.cmd == 'initdb':
        init_db()
    elif args.cmd == 'create-admin':
        if not args.username or not args.password:
            print('Provide --username and --password')
        else:
            create_admin(args.username, args.password)
    else:
        if not os.path.exists('devices.db'):
            init_db()
        app.run(host='0.0.0.0', port=5000, debug=True)
